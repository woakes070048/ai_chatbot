# Copyright (c) 2026, Sanjay Kumar and contributors
# For license information, please see license.txt
"""
DOCX text extraction using python-docx.

Extracts text from paragraphs and tables in Word documents.
"""

from __future__ import annotations

import io


def extract_docx_text(file_bytes: bytes) -> str:
	"""Extract text from a DOCX file.

	Extracts all paragraphs and tables in document order.

	Args:
		file_bytes: Raw DOCX file content.

	Returns:
		Extracted text with paragraphs and tables separated by newlines.
	"""
	try:
		import docx
	except ImportError:
		import frappe

		frappe.throw(
			"python-docx is required for DOCX text extraction. Install it with: pip install python-docx"
		)

	doc = docx.Document(io.BytesIO(file_bytes))
	parts = []

	# Extract paragraphs
	for para in doc.paragraphs:
		text = para.text.strip()
		if text:
			parts.append(text)

	# Extract tables
	for table in doc.tables:
		table_lines = []
		for row in table.rows:
			cells = [cell.text.strip() for cell in row.cells]
			table_lines.append("\t".join(cells))
		if table_lines:
			parts.append("\n".join(table_lines))

	return "\n\n".join(parts)
